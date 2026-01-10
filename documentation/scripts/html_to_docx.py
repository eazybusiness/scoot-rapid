#!/usr/bin/env python3
"""
HTML to DOCX Converter for ScootRapid Documentation
Converts HTML documentation to professional Word documents
"""

import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from PIL import Image
import io

class HTMLToDocxConverter:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup professional document styles"""
        # Title style
        title_style = self.doc.styles['Title']
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(18)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(26, 35, 126)  # ScootRapid blue
        
        # Heading styles
        for i in range(1, 7):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_style.font.name = 'Arial'
            heading_style.font.bold = True
            heading_style.font.color.rgb = RGBColor(57, 73, 171)  # Secondary blue
            
        # Normal style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
    
    def add_heading(self, text, level):
        """Add heading with proper formatting"""
        heading = self.doc.add_heading(text, level)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return heading
    
    def add_paragraph(self, text, style=None):
        """Add paragraph with text"""
        if style:
            para = self.doc.add_paragraph(text, style=style)
        else:
            para = self.doc.add_paragraph(text)
        return para
    
    def add_image(self, image_path, width=Inches(6)):
        """Add image to document"""
        if os.path.exists(image_path):
            para = self.doc.add_paragraph()
            run = para.runs[0] if para.runs else para.add_run()
            run.add_picture(image_path, width=width)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return para
        else:
            print(f"⚠️  Image not found: {image_path}")
            return None
    
    def add_table(self, rows, cols, data):
        """Add table to document"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_data)
                if i == 0:  # Header row
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
        
        return table
    
    def add_code_block(self, code, language=''):
        """Add formatted code block"""
        para = self.doc.add_paragraph()
        run = para.add_run(code)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        
        # Add background color (light gray)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F5F5F5')
        para._p.get_or_add_pPr().insert_element_before(shading, 
            'w:spacing', 'w:ind', 'w:contextualSpacing')
        
        return para
    
    def parse_html_file(self, html_file):
        """Parse HTML file and convert to DOCX"""
        with open(html_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract title
        title = soup.find('title')
        if title:
            self.add_heading(title.text.strip(), 1)
        
        # Process content
        current_section = None
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'pre', 'div']):
            tag = element.name
            
            if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                level = int(tag[1])
                text = element.get_text().strip()
                if text:
                    self.add_heading(text, level)
                    current_section = text
            
            elif tag == 'p':
                text = element.get_text().strip()
                if text:
                    self.add_paragraph(text)
            
            elif tag in ['ul', 'ol']:
                list_items = element.find_all('li')
                for li in list_items:
                    text = li.get_text().strip()
                    if text:
                        self.add_paragraph(f"• {text}")
            
            elif tag == 'table':
                self.process_table(element)
            
            elif tag == 'pre':
                code = element.get_text()
                if code.strip():
                    self.add_code_block(code.strip())
            
            elif tag == 'div' and 'section' in element.get('class', []):
                # Process section divs
                section_title = element.find(['h1', 'h2', 'h3'])
                if section_title:
                    self.add_heading(section_title.get_text().strip(), 2)
    
    def process_table(self, table_element):
        """Process HTML table element"""
        rows = table_element.find_all('tr')
        if not rows:
            return
        
        # Extract table data
        table_data = []
        for row in rows:
            cells = row.find_all(['td', 'th'])
            row_data = [cell.get_text().strip() for cell in cells]
            if row_data:
                table_data.append(row_data)
        
        if table_data:
            self.add_table(len(table_data), len(table_data[0]), table_data)
    
    def add_diagrams_section(self, diagrams_dir):
        """Add diagrams section with images"""
        self.add_heading("Diagramme", 2)
        
        diagram_files = [
            ('Architektur', 'scootrapid_architecture.png'),
            ('Datenbank-Schema', 'scootrapid_database_schema.png'),
            ('API-Fluss', 'scootrapid_api_flow.png'),
            ('Performance-Metriken', 'scootrapid_performance.png'),
            ('Deployment', 'scootrapid_deployment.png'),
        ]
        
        for title, filename in diagram_files:
            image_path = os.path.join(diagrams_dir, filename)
            if os.path.exists(image_path):
                self.add_heading(title, 3)
                self.add_image(image_path, width=Inches(6))
                self.add_paragraph("")  # Add spacing
    
    def add_api_documentation(self, api_commands_file):
        """Add API documentation section"""
        if os.path.exists(api_commands_file):
            self.add_heading("API-Dokumentation", 2)
            
            with open(api_commands_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse markdown file
            lines = content.split('\n')
            current_section = None
            code_block = False
            code_content = []
            
            for line in lines:
                if line.startswith('##'):
                    if current_section and code_content:
                        self.add_code_block('\n'.join(code_content))
                        code_content = []
                    current_section = line.replace('##', '').strip()
                    self.add_heading(current_section, 3)
                    code_block = False
                
                elif line.startswith('```'):
                    if code_block and code_content:
                        self.add_code_block('\n'.join(code_content))
                        code_content = []
                    code_block = not code_block
                
                elif code_block:
                    code_content.append(line)
                
                elif line.startswith('curl') or line.startswith('#'):
                    if code_content:
                        self.add_code_block('\n'.join(code_content))
                        code_content = []
                    self.add_code_block(line)
                
                elif line.strip() and not line.startswith('#'):
                    if code_content:
                        self.add_code_block('\n'.join(code_content))
                        code_content = []
                    self.add_paragraph(line.strip())
            
            # Add remaining code block
            if code_content:
                self.add_code_block('\n'.join(code_content))
    
    def save_document(self, output_file):
        """Save the document"""
        self.doc.save(output_file)
        print(f"✅ Document saved: {output_file}")

def main():
    """Main conversion function"""
    converter = HTMLToDocxConverter()
    
    # Paths
    html_file = "../DOKUMENTATION.html"
    diagrams_dir = "../diagrams"
    api_commands_file = "../api_test_commands.md"
    output_file = "../generated/ScootRapid_Dokumentation.docx"
    
    # Convert HTML
    print("📄 Converting HTML documentation...")
    converter.parse_html_file(html_file)
    
    # Add diagrams
    print("🎨 Adding diagrams...")
    converter.add_diagrams_section(diagrams_dir)
    
    # Add API documentation
    print("🔌 Adding API documentation...")
    converter.add_api_documentation(api_commands_file)
    
    # Save document
    converter.save_document(output_file)
    
    print("✅ ScootRapid documentation converted to DOCX successfully!")

if __name__ == "__main__":
    main()
